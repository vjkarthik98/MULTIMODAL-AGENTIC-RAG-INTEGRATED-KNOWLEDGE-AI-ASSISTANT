import asyncio
import contextvars
import hashlib
import os
import shutil
import subprocess
import tempfile
import time
import uuid
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import structlog
from opentelemetry import trace
from opentelemetry.trace import Status, StatusCode
from prometheus_client import Counter, Histogram
from tenacity import retry, stop_after_attempt, wait_exponential

from app.core.config import settings
from app.ingestion.image_ingest import ingest as image_ingest
from app.ingestion.schema import IngestedDocument

logger = structlog.get_logger(__name__)
tracer = trace.get_tracer(__name__)

# PROMETHEUS METRICS
_ingest_duration = Histogram(
    "document_ingest_duration_seconds",
    "Document ingestion duration by doc type",
    ["doc_type", "status"],
)
_ingest_errors = Counter(
    "document_ingest_errors_total",
    "Document ingestion errors by type",
    ["doc_type", "error_type"],
)
_pii_redacted = Counter(
    "pii_entities_redacted_total_doc",
    "PII entities redacted during document ingestion",
    ["entity_type"],
)
_ocr_invocations = Counter(
    "document_ocr_invocations_total",
    "OCR fallback invocations",
    ["doc_type"],
)

# SIZE LIMITS
_SIZE_LIMITS: Dict[str, int] = {
    ".pdf":  settings.MAX_FILE_SIZE_PDF,
    ".docx": settings.MAX_FILE_SIZE_DOCX,
    ".doc":  settings.MAX_FILE_SIZE_DOCX,
    ".xlsx": settings.MAX_FILE_SIZE_XLSX,
    ".xls":  settings.MAX_FILE_SIZE_XLSX,
}

# SEMAPHORE — CAP CONCURRENT DOCUMENT WORKERS
_semaphore = asyncio.Semaphore(5)


# SHA-256 FILE HASH

def _file_hash(file_path: str) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


# SHA-256 CONTENT HASH FOR DEDUP

def _content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


# TABLE ROWS TO MARKDOWN + TEXT

def _table_to_text(rows: Iterable[Iterable[object]]) -> str:
    cleaned = [
        [str(cell or "").strip() for cell in row]
        for row in (rows or [])
        if any(cell for cell in row)
    ]
    if not cleaned:
        return ""
    return "\n".join(" | ".join(row) for row in cleaned)


def _table_to_markdown(rows: List[List[str]]) -> str:
    if not rows:
        return ""
    header    = rows[0]
    separator = ["---"] * len(header)
    body      = rows[1:] if len(rows) > 1 else []
    lines     = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _format_docx_table_rows(rows: List[List[str]]) -> str:
    lines = []
    for row in rows:
        cells = [str(c or "").strip() for c in row]
        if any(cells):
            lines.append("[TABLE: " + " | ".join(cells) + "]")
    return "\n".join(lines)


# QUALITY SCORE

def _quality(text: str) -> float:
    length = len(text)
    if length < 50:
        return 0.2
    if length < 200:
        return 0.5
    return 1.0


# BASE STRUCTURE

def _base_structure(
    doc_id: str,
    session_id: str,
    source_path: str,
    **extra: Any,
) -> Dict[str, Any]:
    return {
        "doc_id":      doc_id,
        "session_id":  session_id,
        "source_path": source_path,
        **extra,
    }


# PII REDACTION

def _redact_pii(text: str) -> Tuple[str, Dict[str, int]]:
    if not settings.PII_DETECTION_ENABLED:
        return text, {}

    entity_counts: Dict[str, int] = {}

    try:
        from presidio_analyzer import AnalyzerEngine
        from presidio_anonymizer import AnonymizerEngine

        entities = getattr(settings, "PII_ENTITIES", [
            "PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER",
            "US_SSN", "CREDIT_CARD", "LOCATION", "IP_ADDRESS",
        ])

        analyzer   = AnalyzerEngine()
        anonymizer = AnonymizerEngine()
        results    = analyzer.analyze(text=text, entities=entities, language="en")

        for r in results:
            entity_counts[r.entity_type] = entity_counts.get(r.entity_type, 0) + 1

        if results:
            text = anonymizer.anonymize(text=text, analyzer_results=results).text

    except ImportError:
        logger.warning("presidio_not_installed")
    except Exception as exc:
        logger.warning("pii_redaction_failed", error=str(exc))

    return text, entity_counts


# PASSWORD PROTECTION DETECTION — PDF

def _is_pdf_encrypted(file_path: str) -> bool:
    try:
        import fitz
        doc = fitz.open(file_path)
        encrypted = doc.is_encrypted
        doc.close()
        return encrypted
    except Exception:
        return False


# PASSWORD PROTECTION DETECTION — DOCX

def _is_docx_encrypted(file_path: str) -> bool:
    try:
        import zipfile
        with zipfile.ZipFile(file_path, "r") as z:
            names = z.namelist()
            if "EncryptedPackage" in names or "EncryptionInfo" in names:
                return True
        return False
    except zipfile.BadZipFile:
        return True
    except Exception:
        return False


# PDF XREF REPAIR VIA QPDF

def _repair_pdf(file_path: str) -> str:
    try:
        repaired = file_path + ".repaired.pdf"
        result   = subprocess.run(
            ["qpdf", "--replace-input", file_path, "--", repaired],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if os.path.exists(repaired):
            return repaired
    except Exception as exc:
        logger.warning("pdf_repair_failed", error=str(exc))
    return file_path


# PDF JAVASCRIPT DETECTION AND STRIP WARNING

def _check_pdf_javascript(file_path: str) -> bool:
    try:
        import fitz
        doc      = fitz.open(file_path)
        has_js   = False
        for page in doc:
            annots = page.annots()
            if annots:
                for a in annots:
                    if a.info.get("content", "").lower().find("javascript") != -1:
                        has_js = True
                        break
        doc.close()
        return has_js
    except Exception:
        return False


# PDF ROTATION DETECTION AND CORRECTION

def _get_page_rotation(page: Any) -> int:
    try:
        return page.rotation
    except Exception:
        return 0


# MULTI-COLUMN READING ORDER DETECTION

def _correct_reading_order(text: str) -> str:
    # BASIC HEURISTIC — JOIN SHORT LINES THAT LOOK LIKE COLUMN SPLITS
    lines  = text.split("\n")
    output = []
    buffer = ""
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buffer:
                output.append(buffer.strip())
                buffer = ""
            output.append("")
        elif len(stripped) < 40 and not stripped.endswith("."):
            buffer += " " + stripped
        else:
            if buffer:
                output.append(buffer.strip())
                buffer = ""
            output.append(stripped)
    if buffer:
        output.append(buffer.strip())
    return "\n".join(output)


# PDF TEXT DENSITY CHECK FOR OCR FALLBACK

def _text_density(text: str, page_area: float) -> float:
    if page_area <= 0:
        return 0.0
    return len(text.strip()) / page_area


# OCR A SINGLE PAGE IMAGE

def _ocr_page_image(pix: Any, page_num: int) -> Tuple[str, float]:
    try:
        import pytesseract
        from PIL import Image

        img       = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        ocr_text  = (pytesseract.image_to_string(img) or "").strip()

        # PAGE-LEVEL OCR CONFIDENCE
        data       = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        confs      = [int(c) for c in data["conf"] if str(c).lstrip("-").isdigit() and int(c) >= 0]
        confidence = round(sum(confs) / max(len(confs), 1) / 100.0, 3) if confs else 0.5

        return ocr_text, confidence
    except Exception as exc:
        logger.warning("ocr_page_failed", page=page_num, error=str(exc))
        return "", 0.0


# DETECT PDF/A FORMAT

def _is_pdfa(file_path: str) -> bool:
    try:
        with open(file_path, "rb") as f:
            header = f.read(8192).decode("latin-1", errors="ignore")
        return "PDF/A" in header or "pdfa" in header.lower()
    except Exception:
        return False


# DETECT XFA DYNAMIC FORM

def _has_xfa(file_path: str) -> bool:
    try:
        import fitz
        doc    = fitz.open(file_path)
        result = False
        for page in doc:
            if "XFA" in (page.get_text() or ""):
                result = True
                break
        doc.close()
        return result
    except Exception:
        return False


# PDF PROCESSOR

def _process_pdf(
    file_path: str,
    doc_id: str,
    session_id: str,
    source_name: str,
    source_path: str,
    file_size: int = 0,
) -> List[IngestedDocument]:

    import fitz
    import pdfplumber

    documents: List[IngestedDocument] = []

    # SECURITY CHECKS
    if _is_pdf_encrypted(file_path):
        logger.warning("pdf_password_protected_skipped", file=source_name)
        raise ValueError(f"PASSWORD_PROTECTED_PDF: {source_name}")

    has_js = _check_pdf_javascript(file_path)
    if has_js:
        logger.warning("pdf_javascript_detected", file=source_name)

    is_pdfa = _is_pdfa(file_path)
    if is_pdfa:
        logger.info("pdf_archive_format_detected", file=source_name)

    has_xfa = _has_xfa(file_path)
    if has_xfa:
        logger.warning("pdf_xfa_dynamic_form_detected", file=source_name)

    from app.utils.paths import resolved_images_dir
    image_dir = resolved_images_dir() / doc_id
    image_dir.mkdir(parents=True, exist_ok=True)

    header_footer_counts: Dict[str, int] = {}
    active_path = file_path

    try:
        pdf = fitz.open(active_path)
    except Exception:
        # ATTEMPT XREF REPAIR
        logger.warning("pdf_corrupt_attempting_repair", file=source_name)
        active_path = _repair_pdf(file_path)
        try:
            pdf = fitz.open(active_path)
        except Exception as exc:
            raise ValueError(f"PDF_CORRUPT_UNRECOVERABLE: {exc}")

    total_pages = len(pdf)

    for i, page in enumerate(pdf, start=1):
        page_text = ""
        ocr_conf  = 1.0

        # AUTO-ROTATION CORRECTION
        rotation = _get_page_rotation(page)
        if rotation not in (0, 360):
            page.set_rotation(0)

        raw_text = (page.get_text() or "").strip()

        # PRE-03/POST-02: Sanitize extracted PDF text before chunking.
        # pypdf/fitz extracts all text including white-font hidden injection.
        # sanitize() strips injection payloads non-blocking so ingestion continues.
        if raw_text:
            try:
                from app.guardrails.input_guard import sanitize as _guard_sanitize
                raw_text = _guard_sanitize(raw_text, surface="pdf_ingest")
            except Exception:
                pass

        # PAGE AREA FOR DENSITY CHECK
        rect      = page.rect
        page_area = max(rect.width * rect.height, 1)
        density   = _text_density(raw_text, page_area)

        # OCR FALLBACK — only when fitz found no text at all (scanned/image PDF).
        # Low density alone is not a reliable signal: a sparse text page on a
        # large canvas still has real text. Replacing non-empty fitz output with
        # a failed OCR result (empty string) caused NO_CONTENT_EXTRACTED on
        # normal PDFs when tesseract is unavailable.
        if not raw_text:
            _ocr_invocations.labels(doc_type="pdf").inc()
            try:
                pix = page.get_pixmap(dpi=200)
                ocr_result, ocr_conf = _ocr_page_image(pix, i)
                if ocr_result:
                    raw_text = ocr_result
            except Exception as exc:
                logger.warning("pdf_ocr_fallback_failed", page=i, error=str(exc))
        elif density < 0.01:
            # Low density but text exists — OCR as a supplemental pass; keep
            # whichever result is longer (fitz text may be complete, OCR adds
            # nothing when tesseract is absent).
            _ocr_invocations.labels(doc_type="pdf").inc()
            try:
                pix = page.get_pixmap(dpi=200)
                ocr_result, ocr_conf = _ocr_page_image(pix, i)
                if ocr_result and len(ocr_result) > len(raw_text):
                    raw_text = ocr_result
            except Exception as exc:
                logger.warning("pdf_ocr_fallback_failed", page=i, error=str(exc))

        page_text = raw_text.strip()

        # COLLECT HEADER/FOOTER CANDIDATES
        lines = page_text.split("\n")
        if len(lines) > 2:
            header_footer_counts[lines[0]]  = header_footer_counts.get(lines[0], 0) + 1
            header_footer_counts[lines[-1]] = header_footer_counts.get(lines[-1], 0) + 1

        # READING ORDER CORRECTION
        page_text = _correct_reading_order(page_text)

        if page_text:
            page_text, pii_counts = _redact_pii(page_text)
            for et, cnt in pii_counts.items():
                _pii_redacted.labels(entity_type=et).inc(cnt)

            documents.append(
                IngestedDocument(
                    text=page_text,
                    modality="text",
                    subtype="page",
                    source_type="pdf",
                    source=source_name,
                    page=i,
                    structure=_base_structure(
                        doc_id, session_id, source_path,
                        page=i,
                        page_number=i,
                        total_pages=total_pages,
                        ocr_confidence=ocr_conf,
                        is_pdfa=is_pdfa,
                        has_xfa=has_xfa,
                        has_javascript=has_js,
                        section_title=None,
                        ingestion_timestamp=time.time(),
                        language="en",
                        file_size_bytes=file_size,
                        content_type="pdf_page",
                        ingestion_time=time.time(),
                    ),
                    extra_metadata={
                        "data_quality_score": _quality(page_text),
                        "importance_score":   _quality(page_text),
                        "modality_weight":    1.0,
                    },
                ).finalize()
            )

        # EMBEDDED IMAGES
        for img_idx, img_ref in enumerate(page.get_images(full=True)):
            try:
                xref      = img_ref[0]
                base      = pdf.extract_image(xref)
                img_bytes = base.get("image")

                if not img_bytes:
                    continue

                img_path = image_dir / f"p{i}_img{img_idx}.png"
                img_path.write_bytes(img_bytes)

                img_docs = image_ingest(str(img_path), session_id)

                for d in img_docs:
                    d.structure.update({
                        "doc_id":       doc_id,
                        "page":         i,
                        "context_text": page_text[:500],
                        "source_path":  source_path,
                    })
                    documents.append(d)

            except Exception as exc:
                logger.warning(
                    "pdf_image_extract_failed",
                    page=i,
                    img_idx=img_idx,
                    error=str(exc),
                )

    # STRIP REPEATED HEADERS / FOOTERS
    repeated = {k for k, v in header_footer_counts.items() if v > 3 and k.strip()}
    if repeated:
        for d in documents:
            if d.modality == "text":
                for r in repeated:
                    d.text = d.text.replace(r, "").strip()

    # TABLES VIA PDFPLUMBER — RUN IN PARALLEL WITH PAGE LOOP
    try:
        with pdfplumber.open(active_path) as pdf_p:
            for i, page in enumerate(pdf_p.pages, start=1):
                for t_idx, table in enumerate(page.extract_tables() or []):
                    rows = [[str(cell or "").strip() for cell in row] for row in table]
                    txt  = _table_to_text(rows)
                    md   = _table_to_markdown(rows)

                    if not txt:
                        continue

                    combined = f"{txt}\n\n{md}"

                    documents.append(
                        IngestedDocument(
                            text=f"[TABLE page {i}]\n{combined}",
                            modality="table",
                            subtype="structured",
                            source_type="pdf",
                            source=source_name,
                            page=i,
                            structure=_base_structure(
                                doc_id, session_id, source_path,
                                page=i,
                                page_number=i,
                                total_pages=total_pages,
                                table_index=t_idx,
                                section_title=None,
                                ingestion_timestamp=time.time(),
                                language="en",
                                file_size_bytes=file_size,
                                content_type="pdf_table",
                                ingestion_time=time.time(),
                            ),
                            extra_metadata={
                                "data_quality_score": _quality(txt),
                                "importance_score":   _quality(txt),
                                "modality_weight":    1.0,
                            },
                        ).finalize()
                    )

        # BOOKMARKS / OUTLINE TREE
        try:
            with pdfplumber.open(active_path) as pdf_p:
                outline = getattr(pdf_p, "outline", None)
                if outline:
                    outline_text = str(outline)[:2000]
                    documents.append(
                        IngestedDocument(
                            text=f"[OUTLINE]\n{outline_text}",
                            modality="text",
                            subtype="heading",
                            source_type="pdf",
                            source=source_name,
                            structure=_base_structure(
                                doc_id, session_id, source_path,
                                content_type="pdf_outline",
                                ingestion_time=time.time(),
                            ),
                            extra_metadata={
                                "data_quality_score": 0.8,
                                "importance_score":   0.8,
                                "modality_weight":    1.0,
                            },
                        ).finalize()
                    )
        except Exception:
            pass

    except Exception as exc:
        logger.warning("pdf_table_extraction_failed", error=str(exc))

    # HYPERLINK EXTRACTION — reuses the already-open fitz handle to avoid parsing the file twice
    try:
        for i, page in enumerate(pdf, start=1):
            links = page.get_links()
            for link in links:
                uri = link.get("uri", "")
                if uri:
                    documents.append(
                        IngestedDocument(
                            text=f"[HYPERLINK page={i}] {uri}",
                            modality="text",
                            subtype="chunk",
                            source_type="pdf",
                            source=source_name,
                            page=i,
                            structure=_base_structure(
                                doc_id, session_id, source_path,
                                page=i,
                                content_type="pdf_hyperlink",
                                ingestion_time=time.time(),
                            ),
                            extra_metadata={
                                "data_quality_score": 0.6,
                                "importance_score":   0.6,
                                "modality_weight":    0.8,
                            },
                        ).finalize()
                    )
    except Exception as exc:
        logger.warning("pdf_hyperlink_extraction_failed", error=str(exc))
    finally:
        pdf.close()

    return documents


# LIBREOFFICE .DOC TO .DOCX CONVERSION

def _convert_doc_to_docx(file_path: str) -> str:
    try:
        out_dir = Path(file_path).parent
        result  = subprocess.run(
            [
                "libreoffice", "--headless", "--convert-to", "docx",
                "--outdir", str(out_dir), file_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        converted = str(out_dir / (Path(file_path).stem + ".docx"))
        if os.path.exists(converted):
            logger.info("libreoffice_conversion_success", output=converted)
            return converted
        raise RuntimeError(f"LIBREOFFICE_CONVERSION_FAILED: {result.stderr[:300]}")
    except Exception as exc:
        raise RuntimeError(f"DOC_CONVERSION_ERROR: {exc}")


# DOCX ZIP REPAIR ATTEMPT

def _repair_docx(file_path: str) -> str:
    try:
        import zipfile
        repaired = file_path + ".repaired.docx"
        with zipfile.ZipFile(file_path, "r") as zin:
            with zipfile.ZipFile(repaired, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    try:
                        zout.writestr(item, zin.read(item.filename))
                    except Exception:
                        pass
        return repaired
    except Exception as exc:
        logger.warning("docx_repair_failed", error=str(exc))
        return file_path


# MACRO DETECTION IN .DOCM

def _has_macros(file_path: str) -> bool:
    return Path(file_path).suffix.lower() == ".docm"


# DOCX PROCESSOR

# EMBEDDED IMAGE HELPER — write image bytes to a temp file and route through
# image_ingest so we get caption + OCR (text_collection) + SigLIP (vision_collection)
# all linked back to the parent docx/xlsx via parent_* fields.

def _ingest_embedded_image_bytes(
    blob:            bytes,
    ext_hint:        str,
    session_id:      str,
    parent_doc_id:   str,
    parent_modality: str,
    parent_source:   str,
    parent_page:     Optional[int] = None,
    parent_sheet:    Optional[str] = None,
) -> List[IngestedDocument]:
    if not blob:
        return []

    # Reject suspicious tiny payloads (1x1 trackers etc.).
    if len(blob) < 256:
        return []

    safe_ext = (ext_hint or ".png").lower().strip()
    if safe_ext == ".jpeg":
        safe_ext = ".jpg"
    if safe_ext not in {".png", ".jpg", ".bmp", ".gif", ".webp", ".tiff", ".tif"}:
        safe_ext = ".png"

    from app.utils.paths import resolved_temp_dir
    tmp_dir = resolved_temp_dir() / "embedded_images"
    try:
        tmp_dir.mkdir(parents=True, exist_ok=True)
    except Exception:
        tmp_dir = Path(tempfile.gettempdir())

    digest = hashlib.sha256(blob).hexdigest()
    tmp_path = tmp_dir / f"{digest}{safe_ext}"

    try:
        if not tmp_path.exists():
            tmp_path.write_bytes(blob)
        return image_ingest(
            str(tmp_path),
            session_id,
            parent_doc_id=parent_doc_id,
            parent_modality=parent_modality,
            parent_source=parent_source,
            parent_page=parent_page,
            parent_sheet=parent_sheet,
        ) or []
    except Exception as exc:
        logger.warning(
            "embedded_image_ingest_failed",
            parent_modality=parent_modality,
            parent_source=parent_source,
            sheet=parent_sheet,
            page=parent_page,
            error=str(exc),
        )
        return []


def _process_docx(
    file_path: str,
    doc_id: str,
    session_id: str,
    source_name: str,
    source_path: str,
    file_size: int = 0,
) -> List[IngestedDocument]:

    if _is_docx_encrypted(file_path):
        logger.warning("docx_password_protected_skipped", file=source_name)
        raise ValueError(f"PASSWORD_PROTECTED_DOCX: {source_name}")

    if _has_macros(file_path):
        logger.warning("docx_macro_detected", file=source_name)

    try:
        import docx as python_docx
        doc = python_docx.Document(file_path)
    except Exception:
        # ATTEMPT ZIP REPAIR
        logger.warning("docx_corrupt_attempting_repair", file=source_name)
        repaired = _repair_docx(file_path)
        try:
            import docx as python_docx
            doc = python_docx.Document(repaired)
        except Exception as exc:
            raise ValueError(f"DOCX_CORRUPT_UNRECOVERABLE: {exc}")

    documents: List[IngestedDocument] = []

    # EMPTY DOCUMENT CHECK
    has_content = any((p.text or "").strip() for p in doc.paragraphs) or bool(doc.tables)
    if not has_content:
        raise ValueError("EMPTY_DOCUMENT")

    current_heading: Optional[str] = None

    # HEADING HIERARCHY H1-H9
    def _heading_level(paragraph: Any) -> Optional[int]:
        try:
            style_name = paragraph.style.name if paragraph.style else ""
            if "Heading" in style_name:
                parts = style_name.split()
                for p in parts:
                    if p.isdigit():
                        return int(p)
        except Exception:
            pass
        return None

    # PARAGRAPHS
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if not text:
            continue

        level   = _heading_level(p)
        subtype = "heading" if level else "paragraph"
        if level:
            current_heading = text

        text, pii_counts = _redact_pii(text)
        for et, cnt in pii_counts.items():
            _pii_redacted.labels(entity_type=et).inc(cnt)

        documents.append(
            IngestedDocument(
                text=text,
                modality="text",
                subtype=subtype,
                source_type="word",
                source=source_name,
                structure=_base_structure(
                    doc_id, session_id, source_path,
                    paragraph_index=i,
                    heading_level=level,
                    page_number=None,
                    total_pages=None,
                    section_title=text if subtype == "heading" else current_heading,
                    ingestion_timestamp=time.time(),
                    language="en",
                    file_size_bytes=file_size,
                    content_type="docx_paragraph",
                    ingestion_time=time.time(),
                ),
                extra_metadata={
                    "data_quality_score": _quality(text),
                    "importance_score":   1.2 if subtype == "heading" else _quality(text),
                    "modality_weight":    1.0,
                },
            ).finalize()
        )

    # TABLES
    for t_idx, table in enumerate(doc.tables):
        rows     = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        combined = _format_docx_table_rows(rows)

        if not combined:
            continue

        documents.append(
            IngestedDocument(
                text=combined,
                modality="table",
                subtype="structured",
                source_type="word",
                source=source_name,
                structure=_base_structure(
                    doc_id, session_id, source_path,
                    table_index=t_idx,
                    page_number=None,
                    total_pages=None,
                    section_title=current_heading,
                    ingestion_timestamp=time.time(),
                    language="en",
                    file_size_bytes=file_size,
                    content_type="docx_table",
                    ingestion_time=time.time(),
                ),
                extra_metadata={
                    "data_quality_score": _quality(combined),
                    "importance_score":   _quality(combined),
                    "modality_weight":    1.0,
                },
            ).finalize()
        )

    # COMMENTS
    try:
        from docx.oxml.ns import qn
        comments_part = doc.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        )
        if comments_part:
            comments_xml = comments_part._element
            for comment in comments_xml.findall(qn("w:comment")):
                author    = comment.get(qn("w:author"), "")
                date_str  = comment.get(qn("w:date"), "")
                body_text = " ".join(
                    p.text for p in comment.iter() if hasattr(p, "text") and p.text
                ).strip()
                # PRE-11: Author name is PII — scrub before it enters indexed text.
                try:
                    from app.guardrails.pii import scrub_pii
                    author, _ = scrub_pii(author)
                except Exception:
                    pass
                if body_text:
                    documents.append(
                        IngestedDocument(
                            text=f"[COMMENT by {author} on {date_str}] {body_text}",
                            modality="text",
                            subtype="chunk",
                            source_type="word",
                            source=source_name,
                            structure=_base_structure(
                                doc_id, session_id, source_path,
                                comment_author=author,
                                comment_date=date_str,
                                content_type="docx_comment",
                                ingestion_time=time.time(),
                            ),
                            extra_metadata={
                                "data_quality_score": 0.7,
                                "importance_score":   0.7,
                                "modality_weight":    0.9,
                            },
                        ).finalize()
                    )
    except Exception:
        pass

    # FOOTNOTES AND ENDNOTES
    try:
        for fn_type in ("footnotes", "endnotes"):
            try:
                part = getattr(doc.part, fn_type, None)
                if part:
                    xml_text = part._element.text_content() if hasattr(part._element, "text_content") else ""
                    if xml_text.strip():
                        documents.append(
                            IngestedDocument(
                                text=f"[{fn_type.upper()}] {xml_text[:1000]}",
                                modality="text",
                                subtype="chunk",
                                source_type="word",
                                source=source_name,
                                structure=_base_structure(
                                    doc_id, session_id, source_path,
                                    content_type=f"docx_{fn_type}",
                                    ingestion_time=time.time(),
                                ),
                                extra_metadata={
                                    "data_quality_score": 0.6,
                                    "importance_score":   0.6,
                                    "modality_weight":    0.8,
                                },
                            ).finalize()
                        )
            except Exception:
                pass
    except Exception:
        pass

    # HEADERS AND FOOTERS
    try:
        for section in doc.sections:
            for hf_type, hf_obj in [("header", section.header), ("footer", section.footer)]:
                hf_text = " ".join(p.text.strip() for p in hf_obj.paragraphs if p.text.strip())
                if hf_text:
                    documents.append(
                        IngestedDocument(
                            text=f"[{hf_type.upper()}] {hf_text}",
                            modality="text",
                            subtype="chunk",
                            source_type="word",
                            source=source_name,
                            structure=_base_structure(
                                doc_id, session_id, source_path,
                                content_type=f"docx_{hf_type}",
                                ingestion_time=time.time(),
                            ),
                            extra_metadata={
                                "data_quality_score": 0.5,
                                "importance_score":   0.5,
                                "modality_weight":    0.7,
                            },
                        ).finalize()
                    )
    except Exception:
        pass

    # EMBEDDED IMAGES — route each through image_ingest (caption + OCR + SigLIP).
    try:
        related = getattr(doc.part, "related_parts", {}) or {}
        seen_blobs: set = set()
        for rel_id, part in related.items():
            try:
                ct = getattr(part, "content_type", "") or ""
                if not ct.startswith("image/"):
                    continue
                blob = getattr(part, "blob", None)
                if not blob:
                    continue
                digest = hashlib.sha256(blob).hexdigest()
                if digest in seen_blobs:
                    continue
                seen_blobs.add(digest)
                ext_hint = "." + ct.split("/", 1)[-1].split(";", 1)[0].strip().lower()
                if ext_hint in {".jpeg"}:
                    ext_hint = ".jpg"
                embedded = _ingest_embedded_image_bytes(
                    blob=blob,
                    ext_hint=ext_hint,
                    session_id=session_id,
                    parent_doc_id=doc_id,
                    parent_modality="word",
                    parent_source=source_name,
                    parent_page=None,
                    parent_sheet=None,
                )
                documents.extend(embedded)
            except Exception as exc:
                logger.warning("docx_embedded_image_failed", error=str(exc))
    except Exception as exc:
        logger.warning("docx_embedded_image_scan_failed", error=str(exc))

    return documents


# EXCEL PROCESSOR

def _process_excel(
    file_path: str,
    doc_id: str,
    session_id: str,
    source_name: str,
    source_path: str,
    file_size: int = 0,
    warnings: Optional[List[str]] = None,
) -> List[IngestedDocument]:

    import openpyxl

    if warnings is None:
        warnings = []

    documents: List[IngestedDocument] = []
    ROWS_PER_CHUNK = 500

    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
    except Exception as exc:
        err = str(exc)
        raise ValueError(f"CORRUPTED_FILE: {err}")

    try:
        for sheet_name in wb.sheetnames:
            try:
                ws = wb[sheet_name]

                # PRE-09: Skip hidden rows/columns — openpyxl reads ALL rows
                # including those hidden by the author; attacker can embed injection
                # text in hidden cells invisible to the end user.
                from openpyxl.utils import get_column_letter
                hidden_col_letters = {
                    col_letter
                    for col_letter, cd in ws.column_dimensions.items()
                    if cd.hidden
                }

                all_rows = []
                for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                    rd = ws.row_dimensions.get(row_idx)
                    if rd and rd.hidden:
                        continue
                    cells = [
                        str(c if c is not None else "").strip()
                        for col_idx, c in enumerate(row, start=1)
                        if get_column_letter(col_idx) not in hidden_col_letters
                    ]
                    all_rows.append(cells)

                non_empty = [r for r in all_rows if any(c for c in r)]

                if not non_empty:
                    warnings.append(f"Empty sheet skipped: {sheet_name}")
                    continue

                for batch_start in range(0, len(non_empty), ROWS_PER_CHUNK):
                    batch     = non_empty[batch_start:batch_start + ROWS_PER_CHUNK]
                    row_start = batch_start + 1
                    row_end   = batch_start + len(batch)

                    txt = _table_to_text(batch)
                    if not txt.strip():
                        continue

                    chunk_text = f"[Sheet: {sheet_name}, Rows {row_start}-{row_end}]\n{txt}"
                    chunk_text, pii_counts = _redact_pii(chunk_text)
                    for et, cnt in pii_counts.items():
                        _pii_redacted.labels(entity_type=et).inc(cnt)

                    documents.append(
                        IngestedDocument(
                            text=chunk_text,
                            modality="table",
                            subtype="structured",
                            source_type="excel",
                            source=source_name,
                            structure=_base_structure(
                                doc_id, session_id, source_path,
                                sheet=sheet_name,
                                row_start=row_start,
                                row_end=row_end,
                                page_number=None,
                                total_pages=None,
                                section_title=None,
                                ingestion_timestamp=time.time(),
                                language="en",
                                file_size_bytes=file_size,
                                content_type="excel_sheet",
                                ingestion_time=time.time(),
                            ),
                            extra_metadata={
                                "data_quality_score": _quality(txt),
                                "importance_score":   _quality(txt),
                                "modality_weight":    1.0,
                            },
                        ).finalize()
                    )

                # EMBEDDED IMAGES PER SHEET — route through image_ingest.
                try:
                    sheet_images = list(getattr(ws, "_images", []) or [])
                    for img_obj in sheet_images:
                        try:
                            blob: Optional[bytes] = None
                            ext_hint = ".png"
                            # openpyxl <3.1 uses ._data() callable; >=3.1 stores .ref/.path.
                            data_attr = getattr(img_obj, "_data", None)
                            if callable(data_attr):
                                try:
                                    blob = data_attr()
                                except Exception:
                                    blob = None
                            if not blob:
                                ref = getattr(img_obj, "ref", None) or getattr(img_obj, "path", None)
                                if ref and hasattr(ref, "read"):
                                    try:
                                        ref.seek(0)
                                    except Exception:
                                        pass
                                    blob = ref.read()
                                elif isinstance(ref, (str, os.PathLike)):
                                    try:
                                        with open(ref, "rb") as fh:
                                            blob = fh.read()
                                        ext_hint = os.path.splitext(str(ref))[1] or ext_hint
                                    except Exception:
                                        blob = None
                            fmt = getattr(img_obj, "format", None)
                            if fmt:
                                ext_hint = "." + str(fmt).lower()
                            embedded = _ingest_embedded_image_bytes(
                                blob=blob or b"",
                                ext_hint=ext_hint,
                                session_id=session_id,
                                parent_doc_id=doc_id,
                                parent_modality="excel",
                                parent_source=source_name,
                                parent_page=None,
                                parent_sheet=sheet_name,
                            )
                            documents.extend(embedded)
                        except Exception as exc:
                            logger.warning(
                                "excel_embedded_image_failed",
                                sheet=sheet_name,
                                error=str(exc),
                            )
                except Exception as exc:
                    logger.warning(
                        "excel_image_scan_failed",
                        sheet=sheet_name,
                        error=str(exc),
                    )

                # CHARTS AS TEXT — title, axis titles, series names, data range.
                try:
                    sheet_charts = list(getattr(ws, "_charts", []) or [])
                    for chart_idx, chart in enumerate(sheet_charts):
                        try:
                            def _safe_str(v: Any) -> str:
                                try:
                                    return str(v) if v is not None else ""
                                except Exception:
                                    return ""

                            title_text = ""
                            try:
                                t = chart.title
                                if t is not None:
                                    if hasattr(t, "tx") and t.tx and getattr(t.tx, "rich", None):
                                        runs = []
                                        for p in (t.tx.rich.p or []):
                                            for r in (p.r or []):
                                                runs.append(_safe_str(getattr(r, "t", "")))
                                        title_text = " ".join(r for r in runs if r).strip()
                                    elif hasattr(t, "tx") and t.tx and getattr(t.tx, "strRef", None):
                                        title_text = _safe_str(getattr(t.tx.strRef, "f", ""))
                                    else:
                                        title_text = _safe_str(t)
                            except Exception:
                                title_text = ""

                            chart_type = type(chart).__name__
                            x_axis = ""
                            y_axis = ""
                            try:
                                x_axis = _safe_str(getattr(getattr(chart, "x_axis", None), "title", "") or "")
                                y_axis = _safe_str(getattr(getattr(chart, "y_axis", None), "title", "") or "")
                            except Exception:
                                pass

                            series_names: List[str] = []
                            data_refs: List[str] = []
                            try:
                                for s in getattr(chart, "series", []) or []:
                                    try:
                                        if getattr(s, "tx", None) and getattr(s.tx, "strRef", None):
                                            series_names.append(_safe_str(s.tx.strRef.f))
                                    except Exception:
                                        pass
                                    try:
                                        ref = getattr(getattr(s, "val", None), "numRef", None)
                                        if ref is not None:
                                            data_refs.append(_safe_str(getattr(ref, "f", "")))
                                    except Exception:
                                        pass
                            except Exception:
                                pass

                            chart_parts: List[str] = [f"[Chart: {chart_type} on sheet {sheet_name}]"]
                            if title_text:
                                chart_parts.append(f"Title: {title_text}")
                            if x_axis:
                                chart_parts.append(f"X axis: {x_axis}")
                            if y_axis:
                                chart_parts.append(f"Y axis: {y_axis}")
                            series_names = [s for s in series_names if s]
                            if series_names:
                                chart_parts.append("Series: " + ", ".join(series_names[:12]))
                            data_refs = [r for r in data_refs if r]
                            if data_refs:
                                chart_parts.append("Data: " + "; ".join(data_refs[:6]))

                            chart_text = "\n".join(chart_parts)
                            if len(chart_text) < 16:
                                continue

                            chart_text, pii_counts = _redact_pii(chart_text)
                            for et, cnt in pii_counts.items():
                                _pii_redacted.labels(entity_type=et).inc(cnt)

                            documents.append(
                                IngestedDocument(
                                    text=chart_text,
                                    modality="table",
                                    subtype="chart",
                                    source_type="excel",
                                    source=source_name,
                                    structure=_base_structure(
                                        doc_id, session_id, source_path,
                                        sheet=sheet_name,
                                        chart_index=chart_idx,
                                        chart_type=chart_type,
                                        chart_title=title_text or None,
                                        page_number=None,
                                        total_pages=None,
                                        section_title=title_text or None,
                                        ingestion_timestamp=time.time(),
                                        language="en",
                                        file_size_bytes=file_size,
                                        content_type="excel_chart",
                                        ingestion_time=time.time(),
                                    ),
                                    extra_metadata={
                                        "data_quality_score": _quality(chart_text),
                                        "importance_score":   0.9,
                                        "modality_weight":    1.0,
                                    },
                                ).finalize()
                            )
                        except Exception as exc:
                            logger.warning(
                                "excel_chart_failed",
                                sheet=sheet_name,
                                error=str(exc),
                            )
                except Exception as exc:
                    logger.warning(
                        "excel_chart_scan_failed",
                        sheet=sheet_name,
                        error=str(exc),
                    )

            except ValueError:
                raise
            except Exception as exc:
                logger.warning("excel_sheet_failed", sheet=sheet_name, error=str(exc))
    finally:
        try:
            wb.close()
        except Exception:
            pass

    return documents


# MAIN ASYNC INGEST

async def ingest(file_path: str, session_id: str) -> List[IngestedDocument]:

    if not session_id:
        raise ValueError("SESSION_ID_REQUIRED")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"FILE_NOT_FOUND: {file_path}")

    path = Path(file_path)
    ext  = path.suffix.lower()

    size_limit = _SIZE_LIMITS.get(ext, settings.MAX_FILE_SIZE_MB * 1024 * 1024)
    file_size  = path.stat().st_size

    if file_size == 0:
        raise ValueError("EMPTY_FILE")

    if file_size > size_limit:
        raise ValueError(
            f"FILE_TOO_LARGE: {file_size} bytes exceeds {size_limit} bytes for {ext}"
        )

    doc_type = (
        "pdf"   if ext == ".pdf"          else
        "word"  if ext in {".docx", ".doc"} else
        "excel" if ext in {".xlsx", ".xls"} else
        "unknown"
    )

    with tracer.start_as_current_span("document_ingest") as span:
        span.set_attribute("file.name", path.name)
        span.set_attribute("file.size", file_size)
        span.set_attribute("doc.type", doc_type)
        span.set_attribute("session.id", session_id)

        start = time.time()

        async with _semaphore:
            try:
                doc_id      = str(uuid.uuid4())
                file_hash   = await asyncio.get_event_loop().run_in_executor(
                    None, _file_hash, file_path
                )
                source_name = path.name
                source_path = str(path.resolve())

                logger.info(
                    "doc_ingest_start",
                    file=source_name,
                    ext=ext,
                    size=file_size,
                    doc_type=doc_type,
                    session_id=session_id,
                )

                # .DOC REQUIRES LIBREOFFICE CONVERSION FIRST
                active_path = file_path
                if ext == ".doc":
                    active_path = await asyncio.get_event_loop().run_in_executor(
                        None, _convert_doc_to_docx, file_path
                    )

                doc_warnings: List[str] = []

                # DISPATCH TO CORRECT PROCESSOR
                # Copy the current contextvars (notably _current_user_id from
                # app.utils.paths) into the executor thread so resolved_*_dir()
                # helpers can find the active user. A bare lambda would lose it.
                _ctx = contextvars.copy_context()
                if ext == ".pdf":
                    docs = await asyncio.get_event_loop().run_in_executor(
                        None,
                        lambda: _ctx.run(
                            lambda: _process_pdf(
                                active_path, doc_id, session_id, source_name, source_path,
                                file_size=file_size,
                            )
                        ),
                    )
                elif ext in {".docx", ".doc"}:
                    docs = await asyncio.get_event_loop().run_in_executor(
                        None,
                        lambda: _ctx.run(
                            lambda: _process_docx(
                                active_path, doc_id, session_id, source_name, source_path,
                                file_size=file_size,
                            )
                        ),
                    )
                elif ext in {".xlsx", ".xls"}:
                    docs = await asyncio.get_event_loop().run_in_executor(
                        None,
                        lambda: _ctx.run(
                            lambda: _process_excel(
                                active_path, doc_id, session_id, source_name, source_path,
                                file_size=file_size,
                                warnings=doc_warnings,
                            )
                        ),
                    )
                else:
                    raise ValueError(f"UNSUPPORTED_TYPE: {ext}")

                if not docs:
                    raise ValueError("NO_CONTENT_EXTRACTED")

                # STAMP FILE HASH, CHUNK INDEX, AND TOTAL CHUNKS ON ALL DOCS
                total_chunks = len(docs)
                for idx, d in enumerate(docs):
                    d.structure.setdefault("file_hash", file_hash)
                    d.structure.setdefault("ingestion_time", time.time())
                    d.structure.setdefault("chunk_index", idx)
                    d.structure.setdefault("total_chunks", total_chunks)

                latency = round(time.time() - start, 2)

                _ingest_duration.labels(doc_type=doc_type, status="success").observe(latency)

                span.set_attribute("docs.count", len(docs))
                span.set_status(Status(StatusCode.OK))

                logger.info(
                    "doc_ingest_success",
                    file=source_name,
                    ext=ext,
                    docs=len(docs),
                    latency=latency,
                    session_id=session_id,
                )

                return docs

            except Exception as exc:
                latency    = round(time.time() - start, 2)
                error_type = type(exc).__name__

                _ingest_duration.labels(doc_type=doc_type, status="error").observe(latency)
                _ingest_errors.labels(doc_type=doc_type, error_type=error_type).inc()

                span.set_status(Status(StatusCode.ERROR, str(exc)))
                span.record_exception(exc)

                logger.error(
                    "doc_ingest_failed",
                    file=path.name,
                    ext=ext,
                    session_id=session_id,
                    error=str(exc),
                    error_type=error_type,
                    latency=latency,
                )
                raise


# SYNC WRAPPER

def ingest_sync(file_path: str, session_id: str) -> List[IngestedDocument]:
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as pool:
                future = pool.submit(asyncio.run, ingest(file_path, session_id))
                return future.result()
        return loop.run_until_complete(ingest(file_path, session_id))
    except RuntimeError:
        return asyncio.run(ingest(file_path, session_id))


