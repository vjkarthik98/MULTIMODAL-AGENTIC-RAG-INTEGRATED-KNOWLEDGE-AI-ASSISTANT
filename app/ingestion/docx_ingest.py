"""
DOCX/DOC ingestor — Phase 1 per-modality refactor.

DocxIngestor.extract() → List[RawExtract]   (extraction only; no chunking)
ingest()               → List[IngestedDocument]  (backward-compat; full pipeline)
"""

from __future__ import annotations

import asyncio
import hashlib
import os
import re
import subprocess
import tempfile
import time
import uuid
from pathlib import Path
from typing import Any

from opentelemetry import trace
from opentelemetry.trace import Status, StatusCode
from prometheus_client import Counter, Histogram

from app.core.config import settings
from app.ingestion.base_ingest import BaseIngestor
from app.ingestion.schema import IngestedDocument, RawExtract, UniversalMetadata
from app.utils.logger import get_logger

logger = get_logger(__name__)
tracer = trace.get_tracer(__name__)

_ingest_duration = Histogram(
    "docx_ingest_duration_seconds",
    "DOCX ingestion duration",
    ["status"],
)
_ingest_errors = Counter(
    "docx_ingest_errors_total",
    "DOCX ingestion errors by type",
    ["error_type"],
)
_EXTRACTS_TOTAL = Counter("magik_docx_extracts_total", "Total extracts produced by docx ingestor")
_EXTRACT_ERRORS = Counter("magik_docx_extract_errors_total", "Errors in docx ingestor")

_semaphore = asyncio.Semaphore(5)

_DEFINED_TERM_RE = re.compile(
    r'"([^"]+)"\s+(?:means|shall mean|is defined as)',
    re.IGNORECASE,
)
_CLAUSE_NUM_RE = re.compile(
    r'\b(?:Section|Clause|Article|§)\s*\d+(?:\.\d+)*(?:\([a-z]\))*(?:\([ivx]+\))*\b',
    re.IGNORECASE,
)


# ─── Utilities ────────────────────────────────────────────────────────────────


def _file_hash(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _quality(text: str) -> float:
    l = len(text)
    if l < 50:
        return 0.2
    if l < 200:
        return 0.5
    return 1.0


def _is_docx_encrypted(file_path: str) -> bool:
    try:
        import zipfile

        with zipfile.ZipFile(file_path, "r") as z:
            names = z.namelist()
            if "EncryptedPackage" in names or "EncryptionInfo" in names:
                return True
        return False
    except zipfile.BadZipFile:
        return False
    except Exception:
        return False


def _has_macros(file_path: str) -> bool:
    return Path(file_path).suffix.lower() == ".docm"


def _repair_docx(file_path: str) -> str:
    try:
        import zipfile

        repaired = file_path + ".repaired.docx"
        with zipfile.ZipFile(file_path, "r") as zin:
            with zipfile.ZipFile(repaired, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    try:
                        zout.writestr(item, zin.read(item.filename))
                    except Exception:
                        pass
        return repaired
    except Exception as exc:
        logger.warning("docx_repair_failed", error=str(exc))
        return file_path


def _convert_doc_to_docx(file_path: str) -> str:
    out_dir = Path(file_path).parent
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", str(out_dir), file_path],
        capture_output=True,
        text=True,
        timeout=60,
    )
    converted = str(out_dir / (Path(file_path).stem + ".docx"))
    if os.path.exists(converted):
        return converted
    raise RuntimeError(f"LIBREOFFICE_CONVERSION_FAILED: {result.stderr[:300]}")


def _heading_level(paragraph: Any) -> int | None:
    try:
        style_name = paragraph.style.name if paragraph.style else ""
        if "Heading" in style_name:
            for part in style_name.split():
                if part.isdigit():
                    return int(part)
    except Exception:
        pass
    return None


def _looks_like_heading(text: str, paragraph: Any) -> bool:
    t = text.strip()
    if not t or len(t) > 90:
        return False
    words = t.split()
    if not (1 <= len(words) <= 12):
        return False
    if t[-1] in ".,;":
        return False
    letters = [c for c in t if c.isalpha()]
    if len(letters) < 3:
        return False
    if t == t.upper():
        return True
    try:
        runs = [r for r in paragraph.runs if (r.text or "").strip()]
        if runs and all(bool(r.bold) for r in runs):
            return True
    except Exception:
        pass
    _stop = {
        "of",
        "the",
        "and",
        "for",
        "to",
        "in",
        "on",
        "a",
        "an",
        "or",
        "by",
        "with",
        "at",
        "as",
        "from",
    }
    content = [w for w in words if w[:1].isalpha() and w.lower() not in _stop]
    if 1 <= len(content) <= 9:
        cap = sum(1 for w in content if w[:1].isupper())
        if cap / len(content) >= 0.8:
            return True
    return False


def _iter_body_blocks(doc: Any):
    """Yield ('paragraph', Paragraph, para_idx) / ('table', Table, tbl_idx) in true
    document reading order (python-docx's .paragraphs / .tables lose relative
    order — a document with headings interspersed between tables needs them
    walked together via the body XML so heading attribution stays correct).
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    para_idx = 0
    tbl_idx = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("paragraph", Paragraph(child, doc), para_idx)
            para_idx += 1
        elif child.tag == qn("w:tbl"):
            yield ("table", Table(child, doc), tbl_idx)
            tbl_idx += 1


def _format_table_rows(rows: list[list[str]]) -> str:
    lines = []
    for row in rows:
        cells = [str(c or "").strip() for c in row]
        if any(cells):
            lines.append("[TABLE: " + " | ".join(cells) + "]")
    return "\n".join(lines)


def _pii_scrub(text: str, surface: str) -> str:
    try:
        from app.guardrails.pii import scrub_pii as _gp_scrub

        cleaned, _ = _gp_scrub(text)
        return cleaned
    except Exception:
        return text


def _sanitize_text(text: str, surface: str) -> str:
    try:
        from app.guardrails.input_guard import sanitize as _g

        return _g(text, surface=surface)
    except Exception:
        return text


# ─── Phase 1: DocxIngestor ────────────────────────────────────────────────────


class DocxIngestor(BaseIngestor):
    """Extracts raw content from DOCX/DOC files → List[RawExtract].

    Phase 1 responsibility: file I/O, paragraph/table/comment/image parsing.
    Does NOT chunk. The chunker (Phase 2) handles splitting.
    """

    def health_check(self) -> dict:
        return {
            "modality": "docx",
            "status": "ok",
            "class": self.__class__.__name__,
        }

    async def extract(
        self,
        path: Path,
        metadata: UniversalMetadata,
    ) -> list[RawExtract]:
        source = path.name
        file_path = str(path)
        ext = path.suffix.lower()
        logger.info(
            event="extraction_start", modality="docx", file=str(path), size=path.stat().st_size
        )
        try:
            if _is_docx_encrypted(file_path):
                raise ValueError(f"PASSWORD_PROTECTED_DOCX: {path.name}")

            if _has_macros(file_path):
                logger.warning("docx_macro_detected", file=path.name)

            # .doc → convert first
            active_path = file_path
            if ext == ".doc":
                active_path = await asyncio.get_event_loop().run_in_executor(
                    None, _convert_doc_to_docx, file_path
                )

            try:
                import docx as python_docx

                doc = python_docx.Document(active_path)
            except Exception:
                logger.warning("docx_corrupt_attempting_repair", file=path.name)
                repaired = _repair_docx(active_path)
                try:
                    import docx as python_docx

                    doc = python_docx.Document(repaired)
                except Exception as exc:
                    raise ValueError(f"DOCX_CORRUPT_UNRECOVERABLE: {exc}")

            has_content = any((p.text or "").strip() for p in doc.paragraphs) or bool(doc.tables)
            if not has_content:
                raise ValueError("EMPTY_DOCUMENT")

            extracts: list[RawExtract] = []
            current_heading: str | None = None

            # Paragraphs + tables, walked together in true document order so a
            # table's heading/section attribution reflects the heading it
            # actually sits under, not whichever heading happens to be last in
            # the document (python-docx's .paragraphs/.tables lose that order).
            for block_type, block, idx in _iter_body_blocks(doc):
                if block_type == "paragraph":
                    para = block
                    i = idx
                    text = (para.text or "").strip()
                    if not text:
                        continue

                    level = _heading_level(para)
                    if level is None and _looks_like_heading(text, para):
                        level = 2

                    text = self._sanitize(text, surface="docx_ingest")
                    if not text.strip():
                        continue
                    text = self._scrub_pii(text, surface="docx_ingest")

                    style_name: str | None = None
                    try:
                        style_name = para.style.name if para.style else None
                    except Exception:
                        pass

                    is_bold = False
                    try:
                        runs = [r for r in para.runs if (r.text or "").strip()]
                        if runs:
                            is_bold = all(bool(r.bold) for r in runs)
                    except Exception:
                        pass

                    if level:
                        current_heading = text
                        extracts.append(
                            RawExtract(
                                text=text,
                                extract_type="heading",
                                is_bold=is_bold,
                                style_name=style_name,
                                raw_source_ref=f"docx:{path.name}|para:{i}",
                                extra={
                                    "heading_level": level,
                                    "paragraph_index": i,
                                },
                            )
                        )
                    else:
                        extracts.append(
                            RawExtract(
                                text=text,
                                extract_type="prose",
                                is_bold=is_bold,
                                style_name=style_name,
                                raw_source_ref=f"docx:{path.name}|para:{i}",
                                extra={
                                    "paragraph_index": i,
                                    "section_title": current_heading,
                                    "defined_terms": _DEFINED_TERM_RE.findall(text),
                                    "clause_numbers": _CLAUSE_NUM_RE.findall(text),
                                },
                            )
                        )
                else:
                    table = block
                    t_idx = idx
                    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                    row_texts = []
                    for row in rows:
                        if any(row):
                            combined_row = " | ".join(str(c or "") for c in row)
                            combined_row = self._sanitize(combined_row, surface="docx_table_ingest")
                            combined_row = self._scrub_pii(
                                combined_row, surface="docx_table_ingest"
                            )
                            if combined_row.strip():
                                row_texts.append(combined_row)
                                extracts.append(
                                    RawExtract(
                                        text=combined_row,
                                        extract_type="table_row",
                                        raw_source_ref=f"docx:{path.name}|table:{t_idx}|row:{len(row_texts)}",
                                        extra={
                                            "table_index": t_idx,
                                            "section_title": current_heading,
                                        },
                                    )
                                )

            # Comments
            try:
                from docx.oxml.ns import qn

                comments_part = doc.part.package.part_related_by(
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
                )
                if comments_part:
                    comments_xml = comments_part._element
                    for comment in comments_xml.findall(qn("w:comment")):
                        author = comment.get(qn("w:author"), "")
                        date_str = comment.get(qn("w:date"), "")
                        body_text = " ".join(
                            p.text for p in comment.iter() if hasattr(p, "text") and p.text
                        ).strip()
                        try:
                            from app.guardrails.pii import scrub_pii

                            author, _ = scrub_pii(author)
                        except Exception:
                            pass
                        if body_text:
                            body_text = self._sanitize(body_text, surface="docx_comment_ingest")
                            body_text = self._scrub_pii(body_text, surface="docx_comment_ingest")
                            extracts.append(
                                RawExtract(
                                    text=f"[COMMENT by {author} on {date_str}] {body_text}",
                                    extract_type="comment",
                                    raw_source_ref=f"docx:{path.name}|comment",
                                    extra={"comment_author": author, "comment_date": date_str},
                                )
                            )
            except Exception:
                pass

            # Footnotes / endnotes
            try:
                for fn_type in ("footnotes", "endnotes"):
                    try:
                        part = getattr(doc.part, fn_type, None)
                        if part:
                            xml_text = (
                                part._element.text_content()
                                if hasattr(part._element, "text_content")
                                else ""
                            )
                            xml_text = self._sanitize(xml_text, surface="docx_footnote_ingest")
                            xml_text = self._scrub_pii(xml_text, surface="docx_footnote_ingest")
                            if xml_text.strip():
                                extracts.append(
                                    RawExtract(
                                        text=f"[{fn_type.upper()}] {xml_text[:1000]}",
                                        extract_type="footnote",
                                        raw_source_ref=f"docx:{path.name}|{fn_type}",
                                    )
                                )
                    except Exception:
                        pass
            except Exception:
                pass

            # Headers / Footers
            try:
                for section in doc.sections:
                    for hf_type, hf_obj in [("header", section.header), ("footer", section.footer)]:
                        hf_text = " ".join(
                            p.text.strip() for p in hf_obj.paragraphs if p.text.strip()
                        )
                        hf_text = self._sanitize(hf_text, surface="docx_headerfooter_ingest")
                        hf_text = self._scrub_pii(hf_text, surface="docx_headerfooter_ingest")
                        if hf_text:
                            extracts.append(
                                RawExtract(
                                    text=f"[{hf_type.upper()}] {hf_text}",
                                    extract_type="header_footer",
                                    raw_source_ref=f"docx:{path.name}|{hf_type}",
                                )
                            )
            except Exception:
                pass

            # Embedded images → raw_bytes for chunker to pass to captioner
            try:
                related = getattr(doc.part, "related_parts", {}) or {}
                seen_blobs: set = set()
                img_idx = 0
                for rel_id, part in related.items():
                    try:
                        ct = getattr(part, "content_type", "") or ""
                        if not ct.startswith("image/"):
                            continue
                        blob = getattr(part, "blob", None)
                        if not blob or len(blob) < 256:
                            continue
                        digest = hashlib.sha256(blob).hexdigest()
                        if digest in seen_blobs:
                            continue
                        seen_blobs.add(digest)
                        ext_hint = "." + ct.split("/", 1)[-1].split(";", 1)[0].strip().lower()
                        if ext_hint == ".jpeg":
                            ext_hint = ".jpg"
                        extracts.append(
                            RawExtract(
                                text="",
                                extract_type="image_region",
                                raw_source_ref=f"docx:{path.name}|img:{img_idx}",
                                raw_bytes=blob,
                                extra={"img_ext": ext_hint, "section_title": current_heading},
                            )
                        )
                        img_idx += 1
                    except Exception as exc:
                        logger.warning("docx_embedded_image_failed", error=str(exc))
            except Exception as exc:
                logger.warning("docx_embedded_image_scan_failed", error=str(exc))

            if not extracts:
                raise ValueError("NO_EXTRACTS_PRODUCED")

            _EXTRACTS_TOTAL.inc(len(extracts))
            logger.info(
                event="extraction_complete", modality="docx", file=str(path), extracts=len(extracts)
            )
            return extracts
        except Exception as _exc:
            _EXTRACT_ERRORS.inc()
            logger.error(event="extraction_failed", modality="docx", source=source, error=str(_exc))
            raise


# ─── Backward-compat ingest() — full pipeline ─────────────────────────────────


def _base_structure(doc_id: str, session_id: str, source_path: str, **extra: Any) -> dict[str, Any]:
    return {"doc_id": doc_id, "session_id": session_id, "source_path": source_path, **extra}


def _ingest_embedded_image_bytes(
    blob: bytes,
    ext_hint: str,
    session_id: str,
    parent_doc_id: str,
    parent_modality: str,
    parent_source: str,
    parent_page: int | None = None,
    parent_sheet: str | None = None,
) -> list[IngestedDocument]:
    if not blob or len(blob) < 256:
        return []
    safe_ext = (ext_hint or ".png").lower().strip()
    if safe_ext == ".jpeg":
        safe_ext = ".jpg"
    if safe_ext not in {".png", ".jpg", ".bmp", ".gif", ".webp", ".tiff", ".tif"}:
        safe_ext = ".png"
    from app.utils.paths import resolved_temp_dir

    tmp_dir = resolved_temp_dir() / "embedded_images"
    try:
        tmp_dir.mkdir(parents=True, exist_ok=True)
    except Exception:
        tmp_dir = Path(tempfile.gettempdir())
    digest = hashlib.sha256(blob).hexdigest()
    tmp_path = tmp_dir / f"{digest}{safe_ext}"
    try:
        if not tmp_path.exists():
            tmp_path.write_bytes(blob)
        from app.ingestion.image_ingest import ingest as image_ingest

        return (
            image_ingest(
                str(tmp_path),
                session_id,
                parent_doc_id=parent_doc_id,
                parent_modality=parent_modality,
                parent_source=parent_source,
                parent_page=parent_page,
                parent_sheet=parent_sheet,
            )
            or []
        )
    except Exception as exc:
        logger.warning(
            "embedded_image_ingest_failed",
            parent_source=parent_source,
            sheet=parent_sheet,
            page=parent_page,
            error=str(exc),
        )
        return []


async def ingest(file_path: str, session_id: str) -> list[IngestedDocument]:
    """Backward-compatible entry point. Router imports this until Phase 8."""
    if not session_id:
        raise ValueError("SESSION_ID_REQUIRED")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"FILE_NOT_FOUND: {file_path}")

    file_size = path.stat().st_size
    if file_size == 0:
        raise ValueError("EMPTY_FILE")
    if file_size > settings.MAX_FILE_SIZE_DOCX:
        raise ValueError(f"FILE_TOO_LARGE: {file_size}")

    with tracer.start_as_current_span("docx_ingest") as span:
        span.set_attribute("file.name", path.name)
        span.set_attribute("file.size", file_size)
        span.set_attribute("session.id", session_id)
        start = time.time()

        async with _semaphore:
            try:
                logger.info(
                    "docx_ingest_start", file=path.name, size=file_size, session_id=session_id
                )

                ext = path.suffix.lower()
                doc_id = str(uuid.uuid4())
                source_name = path.name
                source_path_str = str(path.resolve())

                if _is_docx_encrypted(file_path):
                    raise ValueError(f"PASSWORD_PROTECTED_DOCX: {path.name}")
                if _has_macros(file_path):
                    logger.warning("docx_macro_detected", file=path.name)

                active_path = file_path
                if ext == ".doc":
                    active_path = await asyncio.get_event_loop().run_in_executor(
                        None, _convert_doc_to_docx, file_path
                    )

                try:
                    import docx as python_docx

                    doc = python_docx.Document(active_path)
                except Exception:
                    logger.warning("docx_corrupt_attempting_repair", file=path.name)
                    repaired = _repair_docx(active_path)
                    try:
                        import docx as python_docx

                        doc = python_docx.Document(repaired)
                    except Exception as exc:
                        raise ValueError(f"DOCX_CORRUPT_UNRECOVERABLE: {exc}")

                has_content = any((p.text or "").strip() for p in doc.paragraphs) or bool(
                    doc.tables
                )
                if not has_content:
                    raise ValueError("EMPTY_DOCUMENT")

                documents: list[IngestedDocument] = []
                current_heading: str | None = None
                _para_buffer: list[str] = []
                _para_buffer_start_idx = 0
                _chunk_target = min(settings.CHUNK_SIZE, 800)

                def _flush_para_buffer(up_to_idx: int) -> None:
                    if not _para_buffer:
                        return
                    merged = " ".join(_para_buffer).strip()
                    merged = _pii_scrub(merged, surface="docx_para_ingest")
                    if not merged:
                        _para_buffer.clear()
                        return
                    documents.append(
                        IngestedDocument(
                            text=merged,
                            modality="text",
                            subtype="paragraph",
                            source_type="word",
                            source=source_name,
                            structure=_base_structure(
                                doc_id,
                                session_id,
                                source_path_str,
                                paragraph_index=_para_buffer_start_idx,
                                heading_level=None,
                                page_number=None,
                                total_pages=None,
                                section_title=current_heading,
                                ingestion_timestamp=time.time(),
                                language="en",
                                file_size_bytes=file_size,
                                content_type="docx_paragraph",
                                ingestion_time=time.time(),
                            ),
                            extra_metadata={
                                "data_quality_score": _quality(merged),
                                "importance_score": _quality(merged),
                                "modality_weight": 1.0,
                            },
                        ).finalize()
                    )
                    _para_buffer.clear()

                # Paragraphs
                for i, para in enumerate(doc.paragraphs):
                    text = (para.text or "").strip()
                    if not text:
                        continue
                    level = _heading_level(para)
                    if level is None and _looks_like_heading(text, para):
                        level = 2
                    subtype = "heading" if level else "paragraph"
                    try:
                        from app.guardrails.input_guard import sanitize as _guard_sanitize

                        _clean = _guard_sanitize(text, surface="docx_ingest")
                        if _clean != text:
                            logger.warning(
                                "docx_injection_sanitized", file=source_name, paragraph_index=i
                            )
                            text = _clean
                    except Exception as _ge:
                        logger.warning("docx_guardrail_failed", file=source_name, error=str(_ge))
                    if not text.strip():
                        continue

                    if level:
                        _flush_para_buffer(i)
                        current_heading = text
                        heading_text = _pii_scrub(text, surface="docx_para_ingest")
                        if heading_text:
                            documents.append(
                                IngestedDocument(
                                    text=heading_text,
                                    modality="text",
                                    subtype="heading",
                                    source_type="word",
                                    source=source_name,
                                    structure=_base_structure(
                                        doc_id,
                                        session_id,
                                        source_path_str,
                                        paragraph_index=i,
                                        heading_level=level,
                                        page_number=None,
                                        total_pages=None,
                                        section_title=heading_text,
                                        ingestion_timestamp=time.time(),
                                        language="en",
                                        file_size_bytes=file_size,
                                        content_type="docx_paragraph",
                                        ingestion_time=time.time(),
                                    ),
                                    extra_metadata={
                                        "data_quality_score": _quality(heading_text),
                                        "importance_score": 1.2,
                                        "modality_weight": 1.0,
                                    },
                                ).finalize()
                            )
                    else:
                        if not _para_buffer:
                            _para_buffer_start_idx = i
                        _para_buffer.append(text)
                        if sum(len(s) + 1 for s in _para_buffer) >= _chunk_target:
                            _flush_para_buffer(i)

                _flush_para_buffer(len(list(doc.paragraphs)))

                # Tables
                for t_idx, table in enumerate(doc.tables):
                    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                    combined = _sanitize_text(
                        "\n".join(
                            "[TABLE: " + " | ".join(str(c or "") for c in row) + "]"
                            for row in rows
                            if any(row)
                        ),
                        surface="docx_table_ingest",
                    )
                    combined = _pii_scrub(combined, surface="docx_table_ingest")
                    if not combined:
                        continue
                    documents.append(
                        IngestedDocument(
                            text=combined,
                            modality="table",
                            subtype="structured",
                            source_type="word",
                            source=source_name,
                            structure=_base_structure(
                                doc_id,
                                session_id,
                                source_path_str,
                                table_index=t_idx,
                                page_number=None,
                                total_pages=None,
                                section_title=current_heading,
                                ingestion_timestamp=time.time(),
                                language="en",
                                file_size_bytes=file_size,
                                content_type="docx_table",
                                ingestion_time=time.time(),
                            ),
                            extra_metadata={
                                "data_quality_score": _quality(combined),
                                "importance_score": _quality(combined),
                                "modality_weight": 1.0,
                            },
                        ).finalize()
                    )

                # Comments
                try:
                    from docx.oxml.ns import qn

                    comments_part = doc.part.package.part_related_by(
                        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
                    )
                    if comments_part:
                        for comment in comments_part._element.findall(qn("w:comment")):
                            author = comment.get(qn("w:author"), "")
                            date_str = comment.get(qn("w:date"), "")
                            body_text = " ".join(
                                p.text for p in comment.iter() if hasattr(p, "text") and p.text
                            ).strip()
                            try:
                                from app.guardrails.pii import scrub_pii

                                author, _ = scrub_pii(author)
                            except Exception:
                                pass
                            if body_text:
                                body_text = _sanitize_text(body_text, surface="docx_comment_ingest")
                                body_text = _pii_scrub(body_text, surface="docx_comment_ingest")
                                documents.append(
                                    IngestedDocument(
                                        text=f"[COMMENT by {author} on {date_str}] {body_text}",
                                        modality="text",
                                        subtype="chunk",
                                        source_type="word",
                                        source=source_name,
                                        structure=_base_structure(
                                            doc_id,
                                            session_id,
                                            source_path_str,
                                            comment_author=author,
                                            comment_date=date_str,
                                            content_type="docx_comment",
                                            ingestion_time=time.time(),
                                        ),
                                        extra_metadata={
                                            "data_quality_score": 0.7,
                                            "importance_score": 0.7,
                                            "modality_weight": 0.9,
                                        },
                                    ).finalize()
                                )
                except Exception:
                    pass

                # Footnotes / Endnotes
                try:
                    for fn_type in ("footnotes", "endnotes"):
                        try:
                            part = getattr(doc.part, fn_type, None)
                            if part:
                                xml_text = (
                                    part._element.text_content()
                                    if hasattr(part._element, "text_content")
                                    else ""
                                )
                                xml_text = _sanitize_text(xml_text, surface="docx_footnote_ingest")
                                xml_text = _pii_scrub(xml_text, surface="docx_footnote_ingest")
                                if xml_text.strip():
                                    documents.append(
                                        IngestedDocument(
                                            text=f"[{fn_type.upper()}] {xml_text[:1000]}",
                                            modality="text",
                                            subtype="chunk",
                                            source_type="word",
                                            source=source_name,
                                            structure=_base_structure(
                                                doc_id,
                                                session_id,
                                                source_path_str,
                                                content_type=f"docx_{fn_type}",
                                                ingestion_time=time.time(),
                                            ),
                                            extra_metadata={
                                                "data_quality_score": 0.6,
                                                "importance_score": 0.6,
                                                "modality_weight": 0.8,
                                            },
                                        ).finalize()
                                    )
                        except Exception:
                            pass
                except Exception:
                    pass

                # Headers / Footers
                try:
                    for section in doc.sections:
                        for hf_type, hf_obj in [
                            ("header", section.header),
                            ("footer", section.footer),
                        ]:
                            hf_text = " ".join(
                                p.text.strip() for p in hf_obj.paragraphs if p.text.strip()
                            )
                            hf_text = _sanitize_text(hf_text, surface="docx_headerfooter_ingest")
                            hf_text = _pii_scrub(hf_text, surface="docx_headerfooter_ingest")
                            if hf_text:
                                documents.append(
                                    IngestedDocument(
                                        text=f"[{hf_type.upper()}] {hf_text}",
                                        modality="text",
                                        subtype="chunk",
                                        source_type="word",
                                        source=source_name,
                                        structure=_base_structure(
                                            doc_id,
                                            session_id,
                                            source_path_str,
                                            content_type=f"docx_{hf_type}",
                                            ingestion_time=time.time(),
                                        ),
                                        extra_metadata={
                                            "data_quality_score": 0.5,
                                            "importance_score": 0.5,
                                            "modality_weight": 0.7,
                                        },
                                    ).finalize()
                                )
                except Exception:
                    pass

                # Embedded images — collected first, then captioned concurrently
                # so BLIP forward passes overlap instead of running one at a time.
                try:
                    import concurrent.futures as _cf

                    related = getattr(doc.part, "related_parts", {}) or {}
                    seen_blobs: set = set()
                    image_jobs: list[tuple[bytes, str]] = []
                    for _rel_id, part in related.items():
                        try:
                            ct = getattr(part, "content_type", "") or ""
                            if not ct.startswith("image/"):
                                continue
                            blob = getattr(part, "blob", None)
                            if not blob:
                                continue
                            digest = hashlib.sha256(blob).hexdigest()
                            if digest in seen_blobs:
                                continue
                            seen_blobs.add(digest)
                            ext_hint = "." + ct.split("/", 1)[-1].split(";", 1)[0].strip().lower()
                            if ext_hint == ".jpeg":
                                ext_hint = ".jpg"
                            image_jobs.append((blob, ext_hint))
                        except Exception as exc:
                            logger.warning("docx_embedded_image_scan_part_failed", error=str(exc))

                    if image_jobs:

                        def _caption_one(args: tuple[bytes, str]) -> list[IngestedDocument]:
                            blob, ext_hint = args
                            try:
                                return _ingest_embedded_image_bytes(
                                    blob=blob,
                                    ext_hint=ext_hint,
                                    session_id=session_id,
                                    parent_doc_id=doc_id,
                                    parent_modality="word",
                                    parent_source=source_name,
                                )
                            except Exception as exc:
                                logger.warning("docx_embedded_image_failed", error=str(exc))
                                return []

                        max_workers = min(
                            len(image_jobs), getattr(settings, "VIDEO_CAPTION_CONCURRENCY", 3)
                        )
                        with _cf.ThreadPoolExecutor(max_workers=max_workers) as img_pool:
                            for result in img_pool.map(_caption_one, image_jobs):
                                documents.extend(result)
                except Exception as exc:
                    logger.warning("docx_embedded_image_scan_failed", error=str(exc))

                if not documents:
                    raise ValueError("NO_CONTENT_EXTRACTED")

                latency = round(time.time() - start, 2)
                _ingest_duration.labels(status="success").observe(latency)
                span.set_attribute("docs.count", len(documents))
                span.set_status(Status(StatusCode.OK))
                logger.info(
                    "docx_ingest_success",
                    file=path.name,
                    docs=len(documents),
                    latency=latency,
                    session_id=session_id,
                )
                return documents

            except Exception as exc:
                latency = round(time.time() - start, 2)
                error_type = type(exc).__name__
                _ingest_duration.labels(status="error").observe(latency)
                _ingest_errors.labels(error_type=error_type).inc()
                span.set_status(Status(StatusCode.ERROR, str(exc)))
                span.record_exception(exc)
                logger.error(
                    "docx_ingest_failed",
                    file=path.name,
                    session_id=session_id,
                    error=str(exc),
                    error_type=error_type,
                    latency=latency,
                )
                raise


def ingest_sync(file_path: str, session_id: str) -> list[IngestedDocument]:
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures

            with concurrent.futures.ThreadPoolExecutor() as pool:
                future = pool.submit(asyncio.run, ingest(file_path, session_id))
                return future.result()
        return loop.run_until_complete(ingest(file_path, session_id))
    except RuntimeError:
        return asyncio.run(ingest(file_path, session_id))


# ─── Production path: DocxIngestor → DocxChunker ──────────────────────────────


async def ingest_docx_full(file_path: str, session_id: str) -> list[IngestedDocument]:
    """Production DOCX ingestion: DocxIngestor.extract() → DocxChunker.chunk().

    Produces modality='docx' docs with full Phase 1.3/2.3 metadata:
    heading_hierarchy, defined_terms, has_bold_terms, has_italic_terms,
    clause_numbers, finance_entities, char_start, char_end, token_count, etc.
    """
    from app.chunking import chunk_raw_extracts
    from app.ingestion.schema import UniversalMetadata

    if not session_id:
        raise ValueError("SESSION_ID_REQUIRED")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"FILE_NOT_FOUND: {file_path}")

    file_size = path.stat().st_size
    if file_size == 0:
        raise ValueError("EMPTY_FILE")
    if file_size > settings.MAX_FILE_SIZE_DOCX:
        raise ValueError(f"FILE_TOO_LARGE: {file_size}")

    meta = UniversalMetadata(
        source_path=str(path.resolve()),
        modality="docx",
        file_size_bytes=file_size,
        custom_fields={"session_id": session_id},
    )

    ingestor = DocxIngestor()
    extracts = await ingestor.extract(path, meta)

    if not extracts:
        raise ValueError("NO_EXTRACTS_PRODUCED")

    docs = chunk_raw_extracts(extracts, meta, "docx")

    for doc in docs:
        struct = getattr(doc, "structure", None)
        if struct is not None and struct.get("session_id") in (None, "default"):
            struct["session_id"] = session_id

    logger.info(
        event="ingest_docx_full_complete",
        file=path.name,
        extracts=len(extracts),
        docs=len(docs),
        session_id=session_id,
    )
    return docs
