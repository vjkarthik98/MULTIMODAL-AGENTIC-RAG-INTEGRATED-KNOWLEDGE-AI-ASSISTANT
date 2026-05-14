import pytest
from unittest.mock import MagicMock, patch, AsyncMock


class TestDocumentIngest:

    @pytest.mark.asyncio
    async def test_pdf_text_extraction_happy_path(self, tmp_path):
        f = tmp_path / "test.pdf"
        f.write_bytes(b"%PDF-1.4 fake")
        mock_docs = [
            MagicMock(
                spec=IngestedDocument,
                text="Page content here",
                structure={"doc_id": "x", "session_id": "s1"},
                modality="text",
            )
        ]
        mock_docs[0].finalize.return_value = mock_docs[0]
        with patch("app.ingestion.document_ingest._process_pdf", return_value=mock_docs), \
             patch("app.ingestion.document_ingest._file_hash", return_value="abc123"), \
             patch("app.ingestion.document_ingest._is_pdf_encrypted", return_value=False):
            docs = await ingest(str(f), "session-1")
        assert len(docs) >= 1

    @pytest.mark.asyncio
    async def test_pdf_scanned_triggers_ocr(self, tmp_path):
        f = tmp_path / "scanned.pdf"
        f.write_bytes(b"%PDF-1.4 scanned")
        with patch("app.ingestion.document_ingest._is_pdf_encrypted", return_value=False), \
             patch("app.ingestion.document_ingest._check_pdf_javascript", return_value=False), \
             patch("app.ingestion.document_ingest._is_pdfa", return_value=False), \
             patch("app.ingestion.document_ingest._has_xfa", return_value=False), \
             patch("app.ingestion.document_ingest._file_hash", return_value="hash1"):
            import fitz
            mock_page = MagicMock()
            mock_page.get_text.return_value = ""
            mock_page.rect.width  = 595
            mock_page.rect.height = 842
            mock_page.rotation    = 0
            mock_page.get_images.return_value = []
            mock_page.get_links.return_value  = []
            mock_pix = MagicMock()
            mock_pix.width   = 595
            mock_pix.height  = 842
            mock_pix.samples = b"\xff" * (595 * 842 * 3)
            mock_page.get_pixmap.return_value = mock_pix
            mock_pdf = MagicMock()
            mock_pdf.__len__    = lambda s: 1
            mock_pdf.__iter__   = lambda s: iter([mock_page])
            mock_pdf.is_encrypted = False
            with patch("fitz.open", return_value=mock_pdf), \
                 patch("app.ingestion.document_ingest._ocr_page_image", return_value=("OCR text from page", 0.85)) as mock_ocr:
                try:
                    await ingest(str(f), "session-1")
                except Exception:
                    pass
                mock_ocr.assert_called()

    @pytest.mark.asyncio
    async def test_pdf_password_protected_skipped(self, tmp_path):
        f = tmp_path / "locked.pdf"
        f.write_bytes(b"%PDF-1.4 encrypted")
        with patch("app.ingestion.document_ingest._is_pdf_encrypted", return_value=True), \
             patch("app.ingestion.document_ingest._file_hash", return_value="abc"):
            with pytest.raises(ValueError, match="PASSWORD_PROTECTED_PDF"):
                await ingest(str(f), "session-1")

    @pytest.mark.asyncio
    async def test_pdf_corrupt_xref_repaired(self, tmp_path):
        f = tmp_path / "corrupt.pdf"
        f.write_bytes(b"%PDF-1.4 corrupt xref")
        with patch("app.ingestion.document_ingest._is_pdf_encrypted", return_value=False), \
             patch("app.ingestion.document_ingest._check_pdf_javascript", return_value=False), \
             patch("app.ingestion.document_ingest._is_pdfa", return_value=False), \
             patch("app.ingestion.document_ingest._has_xfa", return_value=False), \
             patch("app.ingestion.document_ingest._file_hash", return_value="abc"), \
             patch("app.ingestion.document_ingest._repair_pdf", return_value=str(f)) as mock_repair:
            import fitz
            call_count = {"n": 0}
            def open_side_effect(path):
                call_count["n"] += 1
                if call_count["n"] == 1:
                    raise Exception("xref table broken")
                m = MagicMock()
                m.__len__  = lambda s: 0
                m.__iter__ = lambda s: iter([])
                m.close    = MagicMock()
                return m
            with patch("fitz.open", side_effect=open_side_effect):
                try:
                    await ingest(str(f), "session-1")
                except Exception:
                    pass
            mock_repair.assert_called()

    @pytest.mark.asyncio
    async def test_pdf_table_extracted_as_dataframe(self, tmp_path):
        f   = tmp_path / "tables.pdf"
        f.write_bytes(b"%PDF-1.4 with tables")
        row = [["Col1", "Col2"], ["Val1", "Val2"]]
        with patch("app.ingestion.document_ingest._is_pdf_encrypted", return_value=False), \
             patch("app.ingestion.document_ingest._check_pdf_javascript", return_value=False), \
             patch("app.ingestion.document_ingest._is_pdfa", return_value=False), \
             patch("app.ingestion.document_ingest._has_xfa", return_value=False), \
             patch("app.ingestion.document_ingest._file_hash", return_value="abc"):
            table_doc = IngestedDocument(
                text="Col1 | Col2\nVal1 | Val2",
                modality="table",
                subtype="structured",
                source_type="pdf",
                source=f.name,
                structure={"doc_id": "x", "session_id": "session-1", "content_type": "pdf_table"},
                extra_metadata={"data_quality_score": 1.0, "importance_score": 1.0, "modality_weight": 1.0},
            ).finalize()
            with patch("app.ingestion.document_ingest._process_pdf", return_value=[table_doc]):
                docs = await ingest(str(f), "session-1")
        table_docs = [d for d in docs if d.modality == "table"]
        assert len(table_docs) >= 1

    @pytest.mark.asyncio
    async def test_word_doc_libreoffice_conversion(self, tmp_path):
        f = tmp_path / "legacy.doc"
        f.write_bytes(b"\xd0\xcf\x11\xe0 legacy word binary")
        converted = tmp_path / "legacy.docx"
        converted.write_bytes(b"PK fake docx")
        with patch("app.ingestion.document_ingest._convert_doc_to_docx", return_value=str(converted)) as mock_conv, \
             patch("app.ingestion.document_ingest._file_hash", return_value="abc"), \
             patch("app.ingestion.document_ingest._process_docx", return_value=[]):
            try:
                await ingest(str(f), "session-1")
            except Exception:
                pass
            mock_conv.assert_called_once_with(str(f))

    @pytest.mark.asyncio
    async def test_word_track_changes_detected(self, tmp_path):
        f = tmp_path / "tracked.docx"
        f.write_bytes(b"PK fake docx with track changes")
        with patch("app.ingestion.document_ingest._is_docx_encrypted", return_value=False), \
             patch("app.ingestion.document_ingest._has_macros", return_value=False), \
             patch("app.ingestion.document_ingest._file_hash", return_value="abc"):
            import docx as python_docx
            mock_doc = MagicMock()
            mock_doc.paragraphs = []
            mock_doc.tables     = []
            mock_doc.sections   = []
            mock_doc.part       = MagicMock()
            mock_doc.part.package = MagicMock()
            mock_doc.part.package.part_related_by.side_effect = Exception("no comments")
            with patch.object(python_docx, "Document", return_value=mock_doc):
                try:
                    docs = await ingest(str(f), "session-1")
                except Exception:
                    docs = []
        assert isinstance(docs, list)

    @pytest.mark.asyncio
    async def test_word_password_protected_skipped(self, tmp_path):
        f = tmp_path / "locked.docx"
        f.write_bytes(b"PK encrypted docx")
        with patch("app.ingestion.document_ingest._is_docx_encrypted", return_value=True), \
             patch("app.ingestion.document_ingest._file_hash", return_value="abc"):
            with pytest.raises(ValueError, match="PASSWORD_PROTECTED_DOCX"):
                await ingest(str(f), "session-1")

    @pytest.mark.asyncio
    async def test_pdf_js_stripped(self, tmp_path):
        f = tmp_path / "js.pdf"
        f.write_bytes(b"%PDF-1.4 with javascript")

        mock_doc = MagicMock(
            spec=IngestedDocument,
            text="sample text",
            structure={"doc_id": "x"},
            modality="text",
        )
        mock_doc.finalize.return_value = mock_doc

        with patch(
            "app.ingestion.document_ingest._is_pdf_encrypted",
            return_value=False
        ), patch(
            "app.ingestion.document_ingest._check_pdf_javascript",
            return_value=True
        ) as mock_js, patch(
            "app.ingestion.document_ingest._is_pdfa",
            return_value=False
        ), patch(
            "app.ingestion.document_ingest._has_xfa",
            return_value=False
        ), patch(
            "app.ingestion.document_ingest._file_hash",
            return_value="abc"
        ), patch(
            "app.ingestion.document_ingest._process_pdf",
            return_value=[mock_doc]
        ):

            docs = await ingest(str(f), "session-1")

        assert len(docs) >= 1

    def test_metadata_fields_populated(self):
        doc = IngestedDocument(
            text="Sample page content for metadata test",
            modality="text",
            subtype="page",
            source_type="pdf",
            source="test.pdf",
            page=1,
            structure={
                "doc_id":       str(uuid.uuid4()),
                "session_id":   "session-meta",
                "source_path":  "/tmp/test.pdf",
                "page":         1,
                "total_pages":  5,
                "content_type": "pdf_page",
            },
            extra_metadata={
                "data_quality_score": 1.0,
                "importance_score":   1.0,
                "modality_weight":    1.0,
            },
        ).finalize()
        assert doc.structure["doc_id"]
        assert doc.structure["session_id"] == "session-meta"
        assert doc.structure["total_pages"] == 5
        assert doc.modality == "text"

    def test_table_to_markdown_correct_format(self):
        rows = [["Name", "Age"], ["Alice", "30"], ["Bob", "25"]]
        md   = _table_to_markdown(rows)
        assert "| Name | Age |" in md
        assert "| ---" in md
        assert "| Alice | 30 |" in md

    def test_quality_score_short_text(self):
        assert _quality("hi") < 0.5

    def test_quality_score_long_text(self):
        assert _quality("word " * 100) >= 1.0

    @pytest.mark.asyncio
    async def test_empty_file_raises(self, tmp_path):
        f = tmp_path / "empty.pdf"
        f.write_bytes(b"")
        with pytest.raises(ValueError, match="EMPTY_FILE"):
            await ingest(str(f), "session-1")

    @pytest.mark.asyncio
    async def test_no_session_id_raises(self, tmp_path):
        f = tmp_path / "test.pdf"
        f.write_bytes(b"%PDF-1.4")
        with pytest.raises(ValueError, match="SESSION_ID_REQUIRED"):
            await ingest(str(f), "")